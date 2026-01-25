
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx import Document
from .unit_renderer import nettoyer_unite, formatter_qte, convertir_unite
from docx.oxml.ns import nsdecls

# Génération du document .docx
def generate_docx(recette, nombre_couverts=None):
    document = Document()

    # --- Mise en page ---
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    sectPr = section._sectPr
    cols = parse_xml(
        r'<w:cols xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:num="2" w:space="720"/>'
    )
    sectPr.append(cols)

    styles = document.styles
    normal = styles['Normal']
    normal.font.name = 'Calibri Light'
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(40, 40, 40)

    # --- Titre & Métadonnées ---
    title = document.add_paragraph(recette.titre)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.runs[0]
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0, 102, 204)

    nb_couv = nombre_couverts or recette.nombre_couverts
    document.add_paragraph(f"{recette.description or ''}")
    meta = document.add_paragraph(f"📂 {recette.categorie.nom if recette.categorie else 'Non spécifiée'} | 🍽️ {nb_couv} couverts")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(9)
    meta.runs[0].italic = True

    total = (recette.temps_preparation or 0) + (recette.temps_cuisson or 0)
    tps = document.add_paragraph(
        f"⏱️ Préparation : {recette.temps_preparation or 0} min | "
        f"Cuisson : {recette.temps_cuisson or 0} min | Total : {total} min"
    )
    tps.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tps.runs[0].font.size = Pt(9)
    document.add_paragraph("")

    # --- Image ---
    if recette.image and hasattr(recette.image, 'path'):
        try:
            run = document.add_paragraph().add_run()
            run.add_picture(recette.image.path, width=Inches(3))
            document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            document.add_paragraph("(Image non disponible)").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Ingrédients & calcul nutritionnel ---

    p_ing = document.add_paragraph("🥕 Ingrédients",  style='Heading 1')
    p_ing.runs[0].bold = True
    p_ing.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    total_nutrition = {"kcal": 0, "lipides": 0, "glucides": 0, "proteines": 0}
    for ing in recette.recette_ingredient_units.all().order_by('ordre', 'id'):
        # conversion en float pour éviter les erreurs Decimal / float
        qte_base = float(ing.qte or 0)
        nb_ref = float(recette.nombre_couverts or 1)
        nb_couv_f = float(nombre_couverts or nb_ref)

        qte = qte_base / nb_ref * nb_couv_f
        qte, unite = convertir_unite(qte, ing.unit)
        qte_fmt = formatter_qte(qte)
        unite_fmt = nettoyer_unite(unite, qte)
        if unite == '--':
            document.add_paragraph(f"{ing.description}", style='Heading 4')
        else:
            document.add_paragraph(f"• {qte_fmt} {unite_fmt} {ing.description}", style='List')

    # --- Saut de colonne ---
    document.add_paragraph()._element.append(parse_xml(r'<w:br xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:type="column"/>'))

    # --- Étapes ---
    p_et = document.add_paragraph("👨‍🍳 Étapes", style='Heading 1')

    p_et.runs[0].bold = True
    p_et.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    if recette.etapes:
        for ligne in recette.etapes.splitlines():
            txt = ligne.strip()
            if not txt:
                continue
            if txt.startswith('-'):
                p = document.add_paragraph(txt[1:].strip(), style='Heading 4')
            else:
                p = document.add_paragraph(txt, style='List Bullet')
            p.paragraph_format.space_after = Pt(1)
    else:
        document.add_paragraph("Aucune étape spécifiée.")
    # --- Conseils & Note ---
    document.add_paragraph("")
    document.add_heading("💡 Conseils", level=1)

    for ligne in recette.conseils.splitlines():
        txt = ligne.strip()
        if not txt:
            continue
        if txt.startswith('-'):
            document.add_paragraph(txt[1:].strip(), style='Heading 4')
        else:
            document.add_paragraph(txt, style='Normal')
    # § document.add_paragraph(recette.conseils or "")

    document.add_heading("⭐ Note", level=1)
    stars = "★" * int(recette.note) + "☆" * (5 - int(recette.note))
    p = document.add_paragraph(stars)
    p.runs[0].font.color.rgb = RGBColor(255, 200, 0)

    # --- Ajustement police si trop long ---
    if len(document.paragraphs) > 85:
        for p in document.paragraphs:
            for run in p.runs:
                s = run.font.size.pt if run.font.size else 10
                run.font.size = Pt(max(9, s - 1))

    # --- Pied de page ---
    if not document.sections[0].footer.paragraphs:
        document.sections[0].footer.add_paragraph()
    footer = document.sections[0].footer.paragraphs[0]
    footer.text = "Document généré automatiquement par ABsite.fr"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].italic = True
    footer.runs[0].font.color.rgb = RGBColor(150, 150, 150)

    return document